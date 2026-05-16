"""
routes/template_mgmt.py - 模板管理相关路由 Blueprint
从 app_x1.py 提取，保持原有逻辑不变。
"""

import json
import os
import re
import shutil
import zipfile
from pathlib import Path
from datetime import datetime

from flask import Blueprint, jsonify, request, render_template, send_file, current_app
from flask_login import login_required, current_user

from auth import require_role, require_permission
from database import get_db
from monitor import log_action
from config_loader import load_x1_config
from template_resources import resolve_template_resource, apply_type_default_template, apply_semantic_default_template, set_type_default_template, set_semantic_default_template
from template_rules import resolve_template_rule
from helpers.record_utils import _x_now
from helpers.export_utils import _build_export_payload, _x_select_template
from helpers.settings_utils import _setting_enabled
from payload_normalizer import normalize_project_payload, validate_normalized_project
from helpers.template_utils import (
    _compute_template_home_stats, _verify_stage, _flow_rank_from_scene_code,
    _is_scene_error_code, _compute_home_flow_state, _resolve_scene_state_for_home,
    _human_verify_result, _validate_template_key, _inspect_template_docx,
    _compute_template_scene_state, _summarize_template_detail_rows,
    _compute_operating_room_group_stats
)

template_mgmt_bp = Blueprint('template_mgmt', __name__)

# ---------- 路径配置 ----------
BASE_DIR = Path(__file__).parent.parent
_CFG = load_x1_config(BASE_DIR)
_PATHS = _CFG.get('paths', {})
TEMPLATE_BASE = Path(_CFG.get('template_base', '')).expanduser().resolve()
UPLOADS_DIR = BASE_DIR / _PATHS.get('uploads', 'uploads_x1')
APP_VERSION = _CFG.get('version', 'UNKNOWN_VERSION')


@template_mgmt_bp.route('/admin/templates')
@login_required
@require_permission('admin.templates.view')
def admin_templates():
    """模板管理页面"""
    return render_template('templates.html')




@template_mgmt_bp.route('/admin/api/templates')
@login_required
@require_permission('admin.templates.view')
def admin_api_templates():
    from template_resources import TEMPLATE_MAP, list_registered_template_resources, get_type_template_mapping
    overlay = list_registered_template_resources()
    semantic_options = [
        {'semantic_key': 'pharma.veterinary_gmp_workshop.grade.a', 'type_id': 'veterinary_gmp_workshop'},
        {'semantic_key': 'pharma.veterinary_gmp_workshop.grade.b', 'type_id': 'veterinary_gmp_workshop'},
        {'semantic_key': 'pharma.veterinary_gmp_workshop.grade.c', 'type_id': 'veterinary_gmp_workshop'},
        {'semantic_key': 'pharma.veterinary_gmp_workshop.grade.d', 'type_id': 'veterinary_gmp_workshop'},
        {'semantic_key': 'pharma.gmp_workshop.grade.a', 'type_id': 'gmp_workshop'},
        {'semantic_key': 'pharma.gmp_workshop.grade.b', 'type_id': 'gmp_workshop'},
        {'semantic_key': 'pharma.gmp_workshop.grade.c', 'type_id': 'gmp_workshop'},
        {'semantic_key': 'pharma.gmp_workshop.grade.d', 'type_id': 'gmp_workshop'},
        {'semantic_key': 'food.food_workshop.grade.1', 'type_id': 'food_workshop'},
        {'semantic_key': 'food.food_workshop.grade.2', 'type_id': 'food_workshop'},
        {'semantic_key': 'food.food_workshop.grade.3', 'type_id': 'food_workshop'},
        {'semantic_key': 'food.food_workshop.grade.4', 'type_id': 'food_workshop'},
        {'semantic_key': 'electronics.electronics_workshop.iso.5', 'type_id': 'electronics_workshop'},
        {'semantic_key': 'electronics.electronics_workshop.iso.6', 'type_id': 'electronics_workshop'},
        {'semantic_key': 'electronics.electronics_workshop.iso.7', 'type_id': 'electronics_workshop'},
        {'semantic_key': 'electronics.electronics_workshop.iso.8', 'type_id': 'electronics_workshop'},
        {'semantic_key': 'electronics.electronics_workshop.iso.9', 'type_id': 'electronics_workshop'},
        {'semantic_key': 'hospital.clean_function_room.icu', 'type_id': 'clean_function_room'},
        {'semantic_key': 'hospital.clean_function_room.cssd', 'type_id': 'clean_function_room'},
        {'semantic_key': 'hospital.clean_function_room.dialysis', 'type_id': 'clean_function_room'},
        {'semantic_key': 'hospital.clean_function_room.general', 'type_id': 'clean_function_room'},
        {'semantic_key': 'hospital.operating_room.main.level1', 'type_id': 'operating_room'},
        {'semantic_key': 'hospital.operating_room.main.level2', 'type_id': 'operating_room'},
        {'semantic_key': 'hospital.operating_room.main.level3', 'type_id': 'operating_room'},
        {'semantic_key': 'hospital.operating_room.main.level4', 'type_id': 'operating_room'},
        {'semantic_key': 'hospital.operating_room.eye.level1', 'type_id': 'operating_room'},
        {'semantic_key': 'hospital.operating_room.eye.level2', 'type_id': 'operating_room'},
        {'semantic_key': 'hospital.operating_room.eye.level3', 'type_id': 'operating_room'},
        {'semantic_key': 'hospital.operating_room.eye.level4', 'type_id': 'operating_room'},
        {'semantic_key': 'hospital.operating_room.aux.level1', 'type_id': 'operating_room'},
        {'semantic_key': 'hospital.operating_room.aux.level2', 'type_id': 'operating_room'},
        {'semantic_key': 'hospital.operating_room.aux.level3', 'type_id': 'operating_room'},
        {'semantic_key': 'hospital.operating_room.aux.level4', 'type_id': 'operating_room'},
        {'semantic_key': 'biosafety.animal_room.normal', 'type_id': 'animal_room'},
        {'semantic_key': 'biosafety.animal_room.barrier_main', 'type_id': 'animal_room'},
        {'semantic_key': 'biosafety.animal_room.isolation', 'type_id': 'animal_room'},
        {'semantic_key': 'biosafety.animal_room.barrier_aux.clean_storage', 'type_id': 'animal_room'},
        {'semantic_key': 'biosafety.animal_room.barrier_aux.after_sterilization', 'type_id': 'animal_room'},
        {'semantic_key': 'biosafety.animal_room.barrier_aux.clean_corridor', 'type_id': 'animal_room'},
        {'semantic_key': 'biosafety.animal_room.barrier_aux.dirty_corridor', 'type_id': 'animal_room'},
        {'semantic_key': 'biosafety.animal_room.barrier_aux.buffer', 'type_id': 'animal_room'},
        {'semantic_key': 'biosafety.animal_room.barrier_aux.change_room_2', 'type_id': 'animal_room'},
        {'semantic_key': 'biosafety.animal_room.barrier_aux.cleaning_disinfection', 'type_id': 'animal_room'},
        {'semantic_key': 'biosafety.animal_room.barrier_aux.change_room_1', 'type_id': 'animal_room'},
        {'semantic_key': 'biosafety.bsl.p2', 'type_id': 'bsl'},
        {'semantic_key': 'biosafety.bsl.p3', 'type_id': 'bsl'},
    ]
    result = []
    for tid, cfg in TEMPLATE_MAP.items():
        rel_path = str(cfg.get('path', '')).strip()
        abs_dir = TEMPLATE_BASE / rel_path if rel_path else TEMPLATE_BASE
        builtin_files = list(cfg.get('files', []) or [])
        files = []
        for fname in builtin_files:
            fpath = abs_dir / fname
            files.append({
                'name': fname,
                'exists': fpath.exists(),
                'size': fpath.stat().st_size if fpath.exists() else 0,
            })
        registered_pairs = [(k, v) for k, v in overlay.items() if isinstance(v, dict) and v.get('type_id') == tid]
        registered_pairs = sorted(registered_pairs, key=lambda kv: str((kv[1] or {}).get('last_verified_at', '')), reverse=True)
        mapping = get_type_template_mapping(tid)
        allowed_keys = list(mapping.get('allowed_template_keys', []) or [])
        default_template_key = mapping.get('default_template_key', '') or ''
        last_verified_at = ''
        last_verify_result = ''
        version = ''
        enabled = True
        if registered_pairs:
            _, latest = registered_pairs[0]
            last_verified_at = latest.get('last_verified_at', '')
            last_verify_result = latest.get('last_verify_result', '')
            version = latest.get('version', '')
        if default_template_key:
            default_item = overlay.get(default_template_key) or {}
            enabled = default_item.get('enabled', True)
            if not version:
                version = default_item.get('version', '')
            if not last_verified_at:
                last_verified_at = default_item.get('last_verified_at', '')
            if not last_verify_result:
                last_verify_result = default_item.get('last_verify_result', '')
        enabled_count = sum(1 for _, v in registered_pairs if (v or {}).get('enabled', True))
        missing_count = sum(1 for _, v in registered_pairs if not Path(str((v or {}).get('template_path', ''))).exists())

        stats, home_state = _compute_template_home_stats(tid, overlay)

        row = {
            'id': tid,
            'name': cfg.get('name', tid),
            'domain': cfg.get('domain', ''),
            'path': rel_path,
            'template_count': len(files),
            'valid_count': sum(1 for f in files if f.get('exists')),
            'exists': abs_dir.exists(),
            'registered_keys': len(registered_pairs),
            'candidate_count': len(allowed_keys),
            'enabled_count': enabled_count,
            'missing_count': missing_count,
            'source': '内置模板+注册配置' if registered_pairs else '内置模板',
            'version': version or '—',
            'last_verified_at': last_verified_at,
            'last_verify_result': _human_verify_result(last_verify_result or 'unverified'),
            'default_template_key': default_template_key,
            'enabled': enabled,
            'default_warning': (
                'missing' if default_template_key and (not overlay.get(default_template_key) or not Path(str((overlay.get(default_template_key) or {}).get('template_path', '')).strip()).exists())
                else 'disabled' if default_template_key and overlay.get(default_template_key) and (overlay.get(default_template_key) or {}).get('enabled', True) is False
                else 'unset' if not default_template_key
                else ''
            ),
            'home_status_code': home_state.get('code', 'error'),
            'home_status_text': home_state.get('text', '异常'),
            'home_status_level': home_state.get('level', 'danger'),
            'total': stats.get('total', 1),
            'ready': stats.get('ready', 0),
            'readyBasic': stats.get('readyBasic', 0),
            'pending': stats.get('pending', 0),
            'risk': stats.get('risk', 0),
            'registeredOnly': stats.get('registeredOnly', 0),
            'display_enabled_count': stats.get('display_enabled_count', enabled_count),
        }
        if tid == 'operating_room':
            row['group_stats'] = {
                'main': _compute_operating_room_group_stats('main', overlay),
                'eye': _compute_operating_room_group_stats('eye', overlay),
                'aux': _compute_operating_room_group_stats('aux', overlay),
            }
        result.append(row)
    return jsonify({'templates': result, 'total': len(result), 'overlay_total': len(overlay)})




@template_mgmt_bp.route('/admin/api/template-registry/options')
@login_required
@require_permission('admin.templates.view')
def admin_api_template_registry_options():
    import template_resources as tr
    objects = getattr(tr, 'TEMPLATE_OBJECT_OPTIONS', None)
    if not isinstance(objects, dict) or not objects:
        template_map = getattr(tr, 'TEMPLATE_MAP', {}) or {}
        objects = {}
        for type_id, cfg in template_map.items():
            if not isinstance(cfg, dict):
                continue
            objects[type_id] = {
                'name': cfg.get('name', type_id),
                'label': cfg.get('name', type_id),
                'domain': cfg.get('domain', ''),
                'path': cfg.get('path', ''),
                'keyBase': type_id,
            }
    else:
        normalized = {}
        for type_id, cfg in objects.items():
            if not isinstance(cfg, dict):
                continue
            normalized[type_id] = {
                **cfg,
                'label': cfg.get('label') or cfg.get('name') or type_id,
                'keyBase': cfg.get('keyBase') or type_id,
            }
        objects = normalized
    semantic_options = [
        {'semantic_key': 'pharma.veterinary_gmp_workshop.grade.a', 'type_id': 'veterinary_gmp_workshop', 'label': '兽药GMP车间 / A级'},
        {'semantic_key': 'pharma.veterinary_gmp_workshop.grade.b', 'type_id': 'veterinary_gmp_workshop', 'label': '兽药GMP车间 / B级'},
        {'semantic_key': 'pharma.veterinary_gmp_workshop.grade.c', 'type_id': 'veterinary_gmp_workshop', 'label': '兽药GMP车间 / C级'},
        {'semantic_key': 'pharma.veterinary_gmp_workshop.grade.d', 'type_id': 'veterinary_gmp_workshop', 'label': '兽药GMP车间 / D级'},
        {'semantic_key': 'pharma.gmp_workshop.grade.a', 'type_id': 'gmp_workshop', 'label': 'GMP车间 / A级'},
        {'semantic_key': 'pharma.gmp_workshop.grade.b', 'type_id': 'gmp_workshop', 'label': 'GMP车间 / B级'},
        {'semantic_key': 'pharma.gmp_workshop.grade.c', 'type_id': 'gmp_workshop', 'label': 'GMP车间 / C级'},
        {'semantic_key': 'pharma.gmp_workshop.grade.d', 'type_id': 'gmp_workshop', 'label': 'GMP车间 / D级'},
        {'semantic_key': 'food.food_workshop.grade.1', 'type_id': 'food_workshop', 'label': '食品车间 / Ⅰ级'},
        {'semantic_key': 'food.food_workshop.grade.2', 'type_id': 'food_workshop', 'label': '食品车间 / Ⅱ级'},
        {'semantic_key': 'food.food_workshop.grade.3', 'type_id': 'food_workshop', 'label': '食品车间 / Ⅲ级'},
        {'semantic_key': 'food.food_workshop.grade.4', 'type_id': 'food_workshop', 'label': '食品车间 / Ⅳ级'},
        {'semantic_key': 'electronics.electronics_workshop.iso.5', 'type_id': 'electronics_workshop', 'label': '电子车间 / ISO 5'},
        {'semantic_key': 'electronics.electronics_workshop.iso.6', 'type_id': 'electronics_workshop', 'label': '电子车间 / ISO 6'},
        {'semantic_key': 'electronics.electronics_workshop.iso.7', 'type_id': 'electronics_workshop', 'label': '电子车间 / ISO 7'},
        {'semantic_key': 'electronics.electronics_workshop.iso.8', 'type_id': 'electronics_workshop', 'label': '电子车间 / ISO 8'},
        {'semantic_key': 'electronics.electronics_workshop.iso.9', 'type_id': 'electronics_workshop', 'label': '电子车间 / ISO 9'},
        {'semantic_key': 'hospital.clean_function_room.icu', 'type_id': 'clean_function_room', 'label': '洁净功能房 / ICU'},
        {'semantic_key': 'hospital.clean_function_room.cssd', 'type_id': 'clean_function_room', 'label': '洁净功能房 / 消毒供应中心'},
        {'semantic_key': 'hospital.clean_function_room.dialysis', 'type_id': 'clean_function_room', 'label': '洁净功能房 / 透析室'},
        {'semantic_key': 'hospital.clean_function_room.general', 'type_id': 'clean_function_room', 'label': '洁净功能房 / 通用洁净功能房'},
        {'semantic_key': 'hospital.operating_room.main.level1', 'type_id': 'operating_room', 'label': '手术部 / 百级手术室'},
        {'semantic_key': 'hospital.operating_room.main.level2', 'type_id': 'operating_room', 'label': '手术部 / 千级手术室'},
        {'semantic_key': 'hospital.operating_room.main.level3', 'type_id': 'operating_room', 'label': '手术部 / 万级手术室'},
        {'semantic_key': 'hospital.operating_room.main.level4', 'type_id': 'operating_room', 'label': '手术部 / 十万级手术室'},
        {'semantic_key': 'hospital.operating_room.eye.level1', 'type_id': 'operating_room', 'label': '手术部 / 眼科手术室 百级'},
        {'semantic_key': 'hospital.operating_room.eye.level2', 'type_id': 'operating_room', 'label': '手术部 / 眼科手术室 千级'},
        {'semantic_key': 'hospital.operating_room.eye.level3', 'type_id': 'operating_room', 'label': '手术部 / 眼科手术室 万级'},
        {'semantic_key': 'hospital.operating_room.eye.level4', 'type_id': 'operating_room', 'label': '手术部 / 眼科手术室 十万级'},
        {'semantic_key': 'hospital.operating_room.aux.level1', 'type_id': 'operating_room', 'label': '手术部 / 洁净辅房 局5周6'},
        {'semantic_key': 'hospital.operating_room.aux.level2', 'type_id': 'operating_room', 'label': '手术部 / 洁净辅房 ISO 7'},
        {'semantic_key': 'hospital.operating_room.aux.level3', 'type_id': 'operating_room', 'label': '手术部 / 洁净辅房 ISO 8'},
        {'semantic_key': 'hospital.operating_room.aux.level4', 'type_id': 'operating_room', 'label': '手术部 / 洁净辅房 ISO 8.5'},
        {'semantic_key': 'biosafety.animal_room.normal', 'type_id': 'animal_room', 'label': '动物房 / 普通环境'},
        {'semantic_key': 'biosafety.animal_room.barrier_main', 'type_id': 'animal_room', 'label': '动物房 / 屏障环境主房间'},
        {'semantic_key': 'biosafety.animal_room.isolation', 'type_id': 'animal_room', 'label': '动物房 / 隔离环境'},
        {'semantic_key': 'biosafety.animal_room.barrier_aux.clean_storage', 'type_id': 'animal_room', 'label': '动物房 / 屏障环境洁物储存室'},
        {'semantic_key': 'biosafety.animal_room.barrier_aux.after_sterilization', 'type_id': 'animal_room', 'label': '动物房 / 屏障环境灭菌后室区'},
        {'semantic_key': 'biosafety.animal_room.barrier_aux.clean_corridor', 'type_id': 'animal_room', 'label': '动物房 / 屏障环境洁净走廊'},
        {'semantic_key': 'biosafety.animal_room.barrier_aux.dirty_corridor', 'type_id': 'animal_room', 'label': '动物房 / 屏障环境污物走廊'},
        {'semantic_key': 'biosafety.animal_room.barrier_aux.buffer', 'type_id': 'animal_room', 'label': '动物房 / 屏障环境缓冲间'},
        {'semantic_key': 'biosafety.animal_room.barrier_aux.change_room_2', 'type_id': 'animal_room', 'label': '动物房 / 屏障环境二更'},
        {'semantic_key': 'biosafety.animal_room.barrier_aux.cleaning_disinfection', 'type_id': 'animal_room', 'label': '动物房 / 屏障环境清洗消毒室'},
        {'semantic_key': 'biosafety.animal_room.barrier_aux.change_room_1', 'type_id': 'animal_room', 'label': '动物房 / 屏障环境一更'},
        {'semantic_key': 'biosafety.bsl.p2', 'type_id': 'bsl', 'label': '生物安全实验室 / P2'},
        {'semantic_key': 'biosafety.bsl.p3', 'type_id': 'bsl', 'label': '生物安全实验室 / P3'},
    ]
    semantic_by_type = {}
    for item in semantic_options:
        semantic_by_type.setdefault(item['type_id'], []).append(item)
    for type_id, cfg in objects.items():
        if isinstance(cfg, dict):
            cfg['semanticOptions'] = semantic_by_type.get(type_id, [])
    registry_keys = []
    try:
        registry_keys = sorted(list((tr.list_registered_template_resources() or {}).keys()))
    except Exception:
        registry_keys = []
    return jsonify({'success': True, 'objects': objects, 'template_base': str(TEMPLATE_BASE), 'registry_keys': registry_keys})




@template_mgmt_bp.route('/admin/api/template-registry/register', methods=['POST'])
@login_required
@require_permission('admin.templates.registry.manage')
def admin_api_template_registry_register():
    data = request.get_json(silent=True) or {}
    type_id = str(data.get('type_id', '')).strip()
    template_key = str(data.get('template_key', '')).strip()
    template_name = str(data.get('template_name', '')).strip()
    path_mode = str(data.get('path_mode', 'relative')).strip() or 'relative'
    relative_dir = str(data.get('relative_dir', '')).strip().strip('/')
    relative_path = str(data.get('relative_path', '')).strip().strip('/')
    absolute_path = str(data.get('absolute_path', '')).strip()
    resource_note = str(data.get('resource_note', '')).strip()
    attach_to_type = bool(data.get('attach_to_type', True))
    set_as_default = bool(data.get('set_as_default', False))
    semantic_key = str(data.get('semantic_key', '')).strip()
    attach_to_semantic = bool(data.get('attach_to_semantic', False))
    set_as_semantic_default = bool(data.get('set_as_semantic_default', False))

    if not type_id or not template_key or not template_name:
        return jsonify({'success': False, 'error': 'type_id / template_key / template_name 不能为空'}), 400
    key_error = _validate_template_key(template_key)
    if key_error:
        return jsonify({'success': False, 'error': key_error, 'template_key': template_key}), 400

    if path_mode == 'absolute':
        template_path = absolute_path
    else:
        parts = [p for p in [relative_dir, relative_path or template_name] if p]
        template_path = str(TEMPLATE_BASE / '/'.join(parts)) if parts else str(TEMPLATE_BASE / template_name)

    # 路径穿越防护
    resolved = Path(template_path).resolve()
    template_base_resolved = TEMPLATE_BASE.resolve()
    if not str(resolved).startswith(str(template_base_resolved)):
        return jsonify({'success': False, 'error': '模板路径不允许超出模板基础目录'}), 400

    inspect = _inspect_template_docx(template_path)
    exists = inspect['exists']
    valid = inspect['valid']
    parse_error = inspect['parse_error']

    if exists and not valid:
        return jsonify({
            'success': False,
            'error': parse_error or '模板文件校验失败，请确认是有效 docx',
            'template_key': template_key,
            'template_path': template_path,
            'exists': exists,
            'valid': valid,
            'size': inspect['size'],
        }), 400

    from template_resources import register_template_resource, attach_template_key_to_type, set_type_default_template, attach_template_key_to_semantic, set_semantic_default_template, list_registered_template_resources
    existing = list_registered_template_resources() or {}
    if template_key in existing:
        return jsonify({'success': False, 'error': 'template key 已存在，请使用系统自动生成的唯一 key 或更换 key', 'template_key': template_key}), 400
    payload = {
        'template_path': template_path,
        'template_name': template_name,
        'resource_status': 'confirmed' if (exists and valid) else 'missing',
        'resource_note': resource_note or f'后台注册模板：{type_id}',
        'registered_at': _x_now(),
        'registered_by': getattr(current_user, 'id', 'unknown'),
        'type_id': type_id,
    }
    register_template_resource(template_key, payload)
    mapping = None
    semantic_mapping = None
    if attach_to_type:
        mapping = attach_template_key_to_type(type_id, template_key, updated_by=getattr(current_user, 'id', 'unknown'))
    if set_as_default:
        mapping = set_type_default_template(type_id, template_key, updated_by=getattr(current_user, 'id', 'unknown'), updated_at=_x_now())
    if semantic_key and attach_to_semantic:
        semantic_mapping = attach_template_key_to_semantic(semantic_key, template_key, updated_by=getattr(current_user, 'id', 'unknown'))
    if semantic_key and set_as_semantic_default:
        semantic_mapping = set_semantic_default_template(semantic_key, template_key, updated_by=getattr(current_user, 'id', 'unknown'), updated_at=_x_now())
    log_action(current_user.id if current_user.is_authenticated else 'unknown', '注册模板', template_key, template_name)
    return jsonify({'success': True, 'template_key': template_key, 'template_path': template_path, 'exists': exists, 'valid': valid, 'parse_error': parse_error, 'mapping': mapping, 'semantic_mapping': semantic_mapping, 'attach_to_type': attach_to_type, 'set_as_default': set_as_default, 'semantic_key': semantic_key, 'attach_to_semantic': attach_to_semantic, 'set_as_semantic_default': set_as_semantic_default})






@template_mgmt_bp.route('/admin/api/template-registry/upload-and-register', methods=['POST'])
@login_required
@require_permission('admin.templates.registry.manage')
def admin_api_template_registry_upload_and_register():
    type_id = str(request.form.get('type_id', '')).strip()
    template_key = str(request.form.get('template_key', '')).strip()
    relative_dir = str(request.form.get('relative_dir', '')).strip().strip('/')
    resource_note = str(request.form.get('resource_note', '')).strip()
    version = str(request.form.get('version', 'v1')).strip() or 'v1'
    enabled = str(request.form.get('enabled', 'true')).strip().lower() != 'false'
    attach_to_type = str(request.form.get('attach_to_type', 'true')).strip().lower() != 'false'
    set_as_default = str(request.form.get('set_as_default', 'false')).strip().lower() == 'true'
    semantic_key = str(request.form.get('semantic_key', '')).strip()
    attach_to_semantic = str(request.form.get('attach_to_semantic', 'false')).strip().lower() == 'true'
    set_as_semantic_default = str(request.form.get('set_as_semantic_default', 'false')).strip().lower() == 'true'
    if not type_id or not template_key:
        return jsonify({'success': False, 'error': 'type_id / template_key 不能为空'}), 400
    key_error = _validate_template_key(template_key)
    if key_error:
        return jsonify({'success': False, 'error': key_error, 'template_key': template_key}), 400
    file = request.files.get('file')
    if not file or not file.filename:
        return jsonify({'success': False, 'error': '未选择模板文件'}), 400
    if not file.filename.lower().endswith('.docx'):
        return jsonify({'success': False, 'error': '仅支持 .docx 格式'}), 400

    from template_resources import TEMPLATE_OBJECT_OPTIONS, register_template_resource, attach_template_key_to_type, set_type_default_template, attach_template_key_to_semantic, set_semantic_default_template, list_registered_template_resources
    existing = list_registered_template_resources() or {}
    if template_key in existing:
        return jsonify({'success': False, 'error': 'template key 已存在，请使用系统自动生成的唯一 key 或更换 key', 'template_key': template_key}), 400
    obj = TEMPLATE_OBJECT_OPTIONS.get(type_id, {})
    if not relative_dir:
        relative_dir = str(obj.get('path', '')).strip().strip('/')
    if not relative_dir:
        return jsonify({'success': False, 'error': '无法确定模板保存目录'}), 400

    safe_name = Path(file.filename).name
    target_dir = TEMPLATE_BASE / relative_dir
    # 路径穿越防护
    if not str(target_dir.resolve()).startswith(str(TEMPLATE_BASE.resolve())):
        return jsonify({'success': False, 'error': '相对目录不允许超出模板基础目录'}), 400
    target_dir.mkdir(parents=True, exist_ok=True)
    target_path = target_dir / safe_name
    file.save(str(target_path))

    inspect = _inspect_template_docx(target_path)
    exists = inspect['exists']
    valid = inspect['valid']
    parse_error = inspect['parse_error']

    if exists and not valid:
        try:
            target_path.unlink(missing_ok=True)
        except Exception:
            pass
        return jsonify({
            'success': False,
            'error': parse_error or '上传的模板不是有效 docx，已拒绝入库',
            'template_key': template_key,
            'template_path': str(target_path),
            'exists': exists,
            'valid': valid,
            'size': inspect['size'],
        }), 400

    payload = {
        'template_path': str(target_path),
        'template_name': safe_name,
        'resource_status': 'confirmed' if (exists and valid) else 'missing',
        'resource_note': resource_note or f'后台上传注册模板：{type_id}',
        'registered_at': _x_now(),
        'registered_by': getattr(current_user, 'id', 'unknown'),
        'type_id': type_id,
        'enabled': enabled,
        'version': version,
        'last_verified_at': _x_now(),
        'last_verify_result': 'success' if valid else 'failed',
        'last_verify_error': parse_error,
    }
    register_template_resource(template_key, payload)
    mapping = None
    semantic_mapping = None
    if attach_to_type:
        mapping = attach_template_key_to_type(type_id, template_key, updated_by=getattr(current_user, 'id', 'unknown'))
    if set_as_default:
        mapping = set_type_default_template(type_id, template_key, updated_by=getattr(current_user, 'id', 'unknown'), updated_at=_x_now())
    if semantic_key and attach_to_semantic:
        semantic_mapping = attach_template_key_to_semantic(semantic_key, template_key, updated_by=getattr(current_user, 'id', 'unknown'))
    if semantic_key and set_as_semantic_default:
        semantic_mapping = set_semantic_default_template(semantic_key, template_key, updated_by=getattr(current_user, 'id', 'unknown'), updated_at=_x_now())
    log_action(current_user.id if current_user.is_authenticated else 'unknown', '上传并注册模板', template_key, safe_name)
    return jsonify({
        'success': True,
        'upload_success': exists,
        'register_success': True,
        'template_key': template_key,
        'template_path': str(target_path),
        'template_name': safe_name,
        'exists': exists,
        'valid': valid,
        'enabled': enabled,
        'version': version,
        'parse_error': parse_error,
        'mapping': mapping,
        'semantic_mapping': semantic_mapping,
        'attach_to_type': attach_to_type,
        'set_as_default': set_as_default,
        'semantic_key': semantic_key,
        'attach_to_semantic': attach_to_semantic,
        'set_as_semantic_default': set_as_semantic_default,
    })




@template_mgmt_bp.route('/admin/api/template-registry/smoke-export', methods=['POST'])
@login_required
@require_permission('admin.templates.smoke_export')
def admin_api_template_registry_smoke_export():
    data = request.get_json(silent=True) or {}
    template_key = str(data.get('template_key', '')).strip()
    type_id = str(data.get('type_id', '')).strip()
    if not template_key or not type_id:
        return jsonify({'success': False, 'error': 'template_key / type_id 不能为空'}), 400
    from template_resources import list_registered_template_resources, update_template_resource
    overlay = list_registered_template_resources()
    item = overlay.get(template_key)
    if not item:
        return jsonify({'success': False, 'error': '未找到注册模板'}), 404
    if _verify_stage(item.get('last_verify_result', '')) < 2:
        return jsonify({'success': False, 'error': '请先完成基础验证，通过后才能做试导出验证'}), 400

    template_path = str(item.get('template_path', '')).strip()
    exists = bool(template_path and Path(template_path).exists())
    valid = False
    smoke_error = ''
    if exists:
        try:
            from docx import Document
            Document(template_path)
            valid = True
        except Exception as e:
            smoke_error = str(e)
    else:
        smoke_error = '模板文件不存在'

    current = update_template_resource(template_key, {
        'resource_status': 'confirmed' if exists else 'missing',
        'last_verified_at': _x_now(),
        'last_verify_result': 'smoke_success' if valid else 'smoke_failed',
        'last_verify_error': smoke_error,
    })
    log_action(current_user.id if current_user.is_authenticated else 'unknown', '模板导出验证', template_key, current.get('template_name', ''))
    return jsonify({
        'success': valid,
        'template_key': template_key,
        'type_id': type_id,
        'template_name': current.get('template_name', ''),
        'template_path': current.get('template_path', ''),
        'enabled': current.get('enabled', True),
        'version': current.get('version', 'v1'),
        'last_verified_at': current.get('last_verified_at', ''),
        'last_verify_result': current.get('last_verify_result', ''),
        'last_verify_result_label': _human_verify_result(current.get('last_verify_result', '')),
        'last_verify_error': current.get('last_verify_error', ''),
    })




@template_mgmt_bp.route('/admin/api/template-registry/delete', methods=['POST'])
@login_required
@require_permission('admin.templates.delete')
def admin_api_template_registry_delete():
    data = request.get_json(silent=True) or {}
    template_key = str(data.get('template_key', '')).strip()
    if not template_key:
        return jsonify({'success': False, 'error': 'template_key 不能为空'}), 400
    from template_resources import list_registered_template_resources, _save_registry_overlay, _load_type_mappings, _save_type_mappings, _load_semantic_mappings, _save_semantic_mappings
    overlay = list_registered_template_resources()
    item = overlay.get(template_key)
    if not item:
        return jsonify({'success': False, 'error': '未找到注册模板'}), 404
    removed = overlay.pop(template_key, None)
    _save_registry_overlay(overlay)
    # 清理 type_mappings 中的引用
    tm_warnings = []
    tm = _load_type_mappings()
    for tid, cfg in tm.items():
        if cfg.get('default_template_key') == template_key:
            cfg['default_template_key'] = ''
            tm_warnings.append(f'{tid} 的默认模板已清除')
        allowed = cfg.get('allowed_template_keys', [])
        if template_key in allowed:
            allowed.remove(template_key)
    if tm_warnings:
        _save_type_mappings(tm)
    # 清理 semantic_mappings 中的引用
    sm = _load_semantic_mappings()
    sm_warnings = []
    for sk, cfg in sm.items():
        if cfg.get('default_template_key') == template_key:
            cfg['default_template_key'] = ''
            sm_warnings.append(f'{sk} 的默认模板已清除')
        allowed = cfg.get('allowed_template_keys', [])
        if template_key in allowed:
            allowed.remove(template_key)
    if sm_warnings:
        _save_semantic_mappings(sm)
    log_action(current_user.id if current_user.is_authenticated else 'unknown', '删除模板注册', template_key, (removed or {}).get('template_name', ''))
    warnings = tm_warnings + sm_warnings
    result = {'success': True, 'template_key': template_key, 'template_name': (removed or {}).get('template_name', '')}
    if warnings:
        result['warnings'] = warnings
    return jsonify(result)




@template_mgmt_bp.route('/admin/api/template-registry/toggle', methods=['POST'])
@login_required
@require_permission('admin.templates.toggle')
def admin_api_template_registry_toggle():
    data = request.get_json(silent=True) or {}
    template_key = str(data.get('template_key', '')).strip()
    enabled = bool(data.get('enabled', False))
    if not template_key:
        return jsonify({'success': False, 'error': 'template_key 不能为空'}), 400
    from template_resources import update_template_resource, _load_type_mappings, _load_semantic_mappings
    # 停用时检查是否是默认模板
    warning = ''
    if not enabled:
        tm = _load_type_mappings()
        sm = _load_semantic_mappings()
        affected = []
        for tid, cfg in tm.items():
            if cfg.get('default_template_key') == template_key:
                affected.append(f'对象 {tid} 的默认模板')
        for sk, cfg in sm.items():
            if cfg.get('default_template_key') == template_key:
                affected.append(f'场景 {sk} 的默认模板')
        if affected:
            warning = f'警告：该模板当前是{", ".join(affected)}，停用后导出功能将受影响'
    current = update_template_resource(template_key, {'enabled': enabled})
    log_action(current_user.id if current_user.is_authenticated else 'unknown', '切换模板启停', template_key, f'enabled={enabled}')
    result = {'success': True, 'template_key': template_key, 'enabled': current.get('enabled', True)}
    if warning:
        result['warning'] = warning
    return jsonify(result)




@template_mgmt_bp.route('/admin/api/template-registry/verify', methods=['POST'])
@login_required
@require_permission('admin.templates.verify')
def admin_api_template_registry_verify():
    data = request.get_json(silent=True) or {}
    template_key = str(data.get('template_key', '')).strip()
    if not template_key:
        return jsonify({'success': False, 'error': 'template_key 不能为空'}), 400
    from template_resources import list_registered_template_resources, update_template_resource
    overlay = list_registered_template_resources()
    item = overlay.get(template_key)
    if not item:
        return jsonify({'success': False, 'error': '未找到注册模板'}), 404

    template_path = str(item.get('template_path', '')).strip()
    exists = bool(template_path and Path(template_path).exists())
    valid = False
    verify_error = ''
    if exists:
        try:
            from docx import Document
            Document(template_path)
            valid = True
        except Exception as e:
            verify_error = str(e)
    else:
        verify_error = '模板文件不存在'

    patch = {
        'resource_status': 'confirmed' if exists else 'missing',
        'last_verified_at': _x_now(),
        'last_verify_result': 'success' if valid else 'failed',
        'last_verify_error': verify_error,
    }
    current = update_template_resource(template_key, patch)
    log_action(current_user.id if current_user.is_authenticated else 'unknown', '验证模板', template_key, current.get('template_name', ''))
    return jsonify({
        'success': True,
        'template_key': template_key,
        'template_name': current.get('template_name', ''),
        'template_path': current.get('template_path', ''),
        'resource_status': current.get('resource_status', 'missing'),
        'last_verified_at': current.get('last_verified_at', ''),
        'last_verify_result': current.get('last_verify_result', ''),
        'last_verify_result_label': _human_verify_result(current.get('last_verify_result', '')),
        'last_verify_error': current.get('last_verify_error', ''),
    })




@template_mgmt_bp.route('/admin/api/template-type-mappings/<type_id>')
@login_required
@require_permission('admin.templates.view')
def admin_api_template_type_mapping_detail(type_id):
    from template_resources import get_type_template_mapping, list_registered_template_resources
    mapping = get_type_template_mapping(type_id)
    overlay = list_registered_template_resources()
    items = []
    for key, value in overlay.items():
        if isinstance(value, dict) and value.get('type_id') == type_id:
            items.append({
                'template_key': key,
                'template_name': value.get('template_name', ''),
                'enabled': value.get('enabled', True),
                'version': value.get('version', 'v1'),
                'last_verified_at': value.get('last_verified_at', ''),
                'last_verify_result': value.get('last_verify_result', ''),
                'is_default': key == mapping.get('default_template_key', ''),
                'is_allowed': key in (mapping.get('allowed_template_keys', []) or []),
            })
    return jsonify({'success': True, 'type_id': type_id, 'mapping': mapping, 'items': items})




@template_mgmt_bp.route('/admin/api/template-type-mappings/set-default', methods=['POST'])
@login_required
@require_permission('admin.templates.default.set')
def admin_api_template_type_mapping_set_default():
    data = request.get_json(silent=True) or {}
    type_id = str(data.get('type_id', '')).strip()
    template_key = str(data.get('template_key', '')).strip()
    if not type_id or not template_key:
        return jsonify({'success': False, 'error': 'type_id / template_key 不能为空'}), 400
    from template_resources import list_registered_template_resources, set_type_default_template, update_template_resource
    overlay = list_registered_template_resources()
    current = overlay.get(template_key)
    if not isinstance(current, dict) or current.get('type_id') != type_id:
        return jsonify({'success': False, 'error': '该模板不属于指定检测类型或尚未注册'}), 400
    if _verify_stage(current.get('last_verify_result', '')) < 3:
        return jsonify({'success': False, 'error': '请先完成试导出验证，通过后才能启用为当前模板'}), 400
    # 检查模板文件是否存在
    tpath = current.get('template_path', '')
    if tpath and not Path(tpath).exists():
        return jsonify({'success': False, 'error': f'模板文件不存在：{Path(tpath).name}，请先上传模板文件'}), 400
    update_template_resource(template_key, {'enabled': True})
    mapping = set_type_default_template(type_id, template_key, updated_by=getattr(current_user, 'id', 'unknown'), updated_at=_x_now())
    log_action(current_user.id if current_user.is_authenticated else 'unknown', '设置默认模板', type_id, template_key)
    return jsonify({'success': True, 'type_id': type_id, 'template_key': template_key, 'mapping': mapping})




@template_mgmt_bp.route('/admin/api/template-type-mappings/attach', methods=['POST'])
@login_required
@require_permission('admin.templates.mapping.type_manage')
def admin_api_template_type_mapping_attach():
    data = request.get_json(silent=True) or {}
    type_id = str(data.get('type_id', '')).strip()
    template_key = str(data.get('template_key', '')).strip()
    if not type_id or not template_key:
        return jsonify({'success': False, 'error': 'type_id / template_key 不能为空'}), 400
    from template_resources import list_registered_template_resources, attach_template_key_to_type
    overlay = list_registered_template_resources()
    current = overlay.get(template_key)
    if not isinstance(current, dict) or current.get('type_id') != type_id:
        return jsonify({'success': False, 'error': '该模板不属于指定检测类型或尚未注册'}), 400
    mapping = attach_template_key_to_type(type_id, template_key, updated_by=getattr(current_user, 'id', 'unknown'))
    log_action(current_user.id if current_user.is_authenticated else 'unknown', '加入模板候选', type_id, template_key)
    return jsonify({'success': True, 'type_id': type_id, 'template_key': template_key, 'mapping': mapping})




@template_mgmt_bp.route('/admin/api/template-semantic-mappings/options')
@login_required
@require_permission('admin.templates.view')
def admin_api_template_semantic_mapping_options():
    options = [
        {'semantic_key': 'pharma.veterinary_gmp_workshop.grade.a', 'type_id': 'veterinary_gmp_workshop', 'label': '兽药GMP车间 / A级'},
        {'semantic_key': 'pharma.veterinary_gmp_workshop.grade.b', 'type_id': 'veterinary_gmp_workshop', 'label': '兽药GMP车间 / B级'},
        {'semantic_key': 'pharma.veterinary_gmp_workshop.grade.c', 'type_id': 'veterinary_gmp_workshop', 'label': '兽药GMP车间 / C级'},
        {'semantic_key': 'pharma.veterinary_gmp_workshop.grade.d', 'type_id': 'veterinary_gmp_workshop', 'label': '兽药GMP车间 / D级'},
        {'semantic_key': 'pharma.gmp_workshop.grade.a', 'type_id': 'gmp_workshop', 'label': 'GMP车间 / A级'},
        {'semantic_key': 'pharma.gmp_workshop.grade.b', 'type_id': 'gmp_workshop', 'label': 'GMP车间 / B级'},
        {'semantic_key': 'pharma.gmp_workshop.grade.c', 'type_id': 'gmp_workshop', 'label': 'GMP车间 / C级'},
        {'semantic_key': 'pharma.gmp_workshop.grade.d', 'type_id': 'gmp_workshop', 'label': 'GMP车间 / D级'},
        {'semantic_key': 'food.food_workshop.grade.1', 'type_id': 'food_workshop', 'label': '食品车间 / Ⅰ级'},
        {'semantic_key': 'food.food_workshop.grade.2', 'type_id': 'food_workshop', 'label': '食品车间 / Ⅱ级'},
        {'semantic_key': 'food.food_workshop.grade.3', 'type_id': 'food_workshop', 'label': '食品车间 / Ⅲ级'},
        {'semantic_key': 'food.food_workshop.grade.4', 'type_id': 'food_workshop', 'label': '食品车间 / Ⅳ级'},
        {'semantic_key': 'electronics.electronics_workshop.iso.5', 'type_id': 'electronics_workshop', 'label': '电子车间 / ISO 5'},
        {'semantic_key': 'electronics.electronics_workshop.iso.6', 'type_id': 'electronics_workshop', 'label': '电子车间 / ISO 6'},
        {'semantic_key': 'electronics.electronics_workshop.iso.7', 'type_id': 'electronics_workshop', 'label': '电子车间 / ISO 7'},
        {'semantic_key': 'electronics.electronics_workshop.iso.8', 'type_id': 'electronics_workshop', 'label': '电子车间 / ISO 8'},
        {'semantic_key': 'electronics.electronics_workshop.iso.9', 'type_id': 'electronics_workshop', 'label': '电子车间 / ISO 9'},
        {'semantic_key': 'hospital.clean_function_room.icu', 'type_id': 'clean_function_room', 'label': '洁净功能房 / ICU'},
        {'semantic_key': 'hospital.clean_function_room.cssd', 'type_id': 'clean_function_room', 'label': '洁净功能房 / 消毒供应中心'},
        {'semantic_key': 'hospital.clean_function_room.dialysis', 'type_id': 'clean_function_room', 'label': '洁净功能房 / 透析室'},
        {'semantic_key': 'hospital.clean_function_room.general', 'type_id': 'clean_function_room', 'label': '洁净功能房 / 通用洁净功能房'},
        {'semantic_key': 'hospital.operating_room.main.level1', 'type_id': 'operating_room', 'label': '手术部 / 百级手术室'},
        {'semantic_key': 'hospital.operating_room.main.level2', 'type_id': 'operating_room', 'label': '手术部 / 千级手术室'},
        {'semantic_key': 'hospital.operating_room.main.level3', 'type_id': 'operating_room', 'label': '手术部 / 万级手术室'},
        {'semantic_key': 'hospital.operating_room.main.level4', 'type_id': 'operating_room', 'label': '手术部 / 十万级手术室'},
        {'semantic_key': 'hospital.operating_room.eye.level1', 'type_id': 'operating_room', 'label': '手术部 / 眼科手术室 百级'},
        {'semantic_key': 'hospital.operating_room.eye.level2', 'type_id': 'operating_room', 'label': '手术部 / 眼科手术室 千级'},
        {'semantic_key': 'hospital.operating_room.eye.level3', 'type_id': 'operating_room', 'label': '手术部 / 眼科手术室 万级'},
        {'semantic_key': 'hospital.operating_room.eye.level4', 'type_id': 'operating_room', 'label': '手术部 / 眼科手术室 十万级'},
        {'semantic_key': 'hospital.operating_room.aux.level1', 'type_id': 'operating_room', 'label': '手术部 / 洁净辅房 局5周6'},
        {'semantic_key': 'hospital.operating_room.aux.level2', 'type_id': 'operating_room', 'label': '手术部 / 洁净辅房 ISO 7'},
        {'semantic_key': 'hospital.operating_room.aux.level3', 'type_id': 'operating_room', 'label': '手术部 / 洁净辅房 ISO 8'},
        {'semantic_key': 'hospital.operating_room.aux.level4', 'type_id': 'operating_room', 'label': '手术部 / 洁净辅房 ISO 8.5'},
        {'semantic_key': 'biosafety.animal_room.normal', 'type_id': 'animal_room', 'label': '动物房 / 普通环境'},
        {'semantic_key': 'biosafety.animal_room.barrier_main', 'type_id': 'animal_room', 'label': '动物房 / 屏障环境主房间'},
        {'semantic_key': 'biosafety.animal_room.isolation', 'type_id': 'animal_room', 'label': '动物房 / 隔离环境'},
        {'semantic_key': 'biosafety.animal_room.barrier_aux.clean_storage', 'type_id': 'animal_room', 'label': '动物房 / 屏障环境洁物储存室'},
        {'semantic_key': 'biosafety.animal_room.barrier_aux.after_sterilization', 'type_id': 'animal_room', 'label': '动物房 / 屏障环境灭菌后室区'},
        {'semantic_key': 'biosafety.animal_room.barrier_aux.clean_corridor', 'type_id': 'animal_room', 'label': '动物房 / 屏障环境洁净走廊'},
        {'semantic_key': 'biosafety.animal_room.barrier_aux.dirty_corridor', 'type_id': 'animal_room', 'label': '动物房 / 屏障环境污物走廊'},
        {'semantic_key': 'biosafety.animal_room.barrier_aux.buffer', 'type_id': 'animal_room', 'label': '动物房 / 屏障环境缓冲间'},
        {'semantic_key': 'biosafety.animal_room.barrier_aux.change_room_2', 'type_id': 'animal_room', 'label': '动物房 / 屏障环境二更'},
        {'semantic_key': 'biosafety.animal_room.barrier_aux.cleaning_disinfection', 'type_id': 'animal_room', 'label': '动物房 / 屏障环境清洗消毒室'},
        {'semantic_key': 'biosafety.animal_room.barrier_aux.change_room_1', 'type_id': 'animal_room', 'label': '动物房 / 屏障环境一更'},
        {'semantic_key': 'biosafety.bsl.p2', 'type_id': 'bsl', 'label': '生物安全实验室 / P2'},
        {'semantic_key': 'biosafety.bsl.p3', 'type_id': 'bsl', 'label': '生物安全实验室 / P3'},
    ]
    return jsonify({'success': True, 'options': options})




@template_mgmt_bp.route('/admin/api/template-semantic-mappings/<path:semantic_key>')
@login_required
@require_permission('admin.templates.view')
def admin_api_template_semantic_mapping_detail(semantic_key):
    import re
    if not re.match(r'^[a-zA-Z0-9_.\-/]+$', semantic_key) or len(semantic_key) > 128:
        return jsonify({'success': False, 'error': 'semantic_key 格式无效'}), 400
    from pathlib import Path
    from template_resources import get_semantic_template_mapping, list_registered_template_resources
    mapping = get_semantic_template_mapping(semantic_key)
    overlay = list_registered_template_resources()
    type_id = str(request.args.get('type_id', '')).strip()
    items = []
    allowed_keys = set(mapping.get('allowed_template_keys', []) or [])
    default_key = str(mapping.get('default_template_key', '')).strip()
    default_item = None
    for key, value in overlay.items():
        if not isinstance(value, dict):
            continue
        if type_id and value.get('type_id') != type_id:
            continue
        if (not type_id) and key not in allowed_keys:
            continue
        template_path = str(value.get('template_path', '')).strip()
        path_obj = Path(template_path) if template_path else None
        exists = bool(path_obj and path_obj.exists())
        size = path_obj.stat().st_size if exists else 0
        row = {
            'template_key': key,
            'template_name': value.get('template_name', ''),
            'template_path': template_path,
            'exists': exists,
            'size': size,
            'enabled': value.get('enabled', True),
            'version': value.get('version', 'v1'),
            'last_verified_at': value.get('last_verified_at', ''),
            'last_verify_result': value.get('last_verify_result', ''),
            'last_verify_result_label': _human_verify_result(value.get('last_verify_result', '')),
            'is_default': key == default_key,
            'is_allowed': key in allowed_keys,
        }
        items.append(row)
        if key == default_key:
            default_item = row
    scene_state = _compute_template_scene_state(mapping, default_item)
    return jsonify({
        'success': True,
        'semantic_key': semantic_key,
        'mapping': mapping,
        'items': items,
        'scene_state': scene_state,
    })




@template_mgmt_bp.route('/admin/api/template-semantic-mappings/set-default', methods=['POST'])
@login_required
@require_permission('admin.templates.default.set')
def admin_api_template_semantic_mapping_set_default():
    data = request.get_json(silent=True) or {}
    semantic_key = str(data.get('semantic_key', '')).strip()
    template_key = str(data.get('template_key', '')).strip()
    if not semantic_key or not template_key:
        return jsonify({'success': False, 'error': 'semantic_key / template_key 不能为空'}), 400
    from template_resources import list_registered_template_resources, set_semantic_default_template, update_template_resource
    overlay = list_registered_template_resources()
    current = overlay.get(template_key)
    if not isinstance(current, dict):
        return jsonify({'success': False, 'error': '该模板尚未注册'}), 400
    if _verify_stage(current.get('last_verify_result', '')) < 3:
        return jsonify({'success': False, 'error': '请先完成试导出验证，通过后才能启用为当前模板'}), 400
    # 检查模板文件是否存在
    tpath = current.get('template_path', '')
    if tpath and not Path(tpath).exists():
        return jsonify({'success': False, 'error': f'模板文件不存在：{Path(tpath).name}，请先上传模板文件'}), 400
    # 校验模板 type_id 与 semantic_key 对应的对象类型是否匹配
    template_type_id = current.get('type_id', '')
    # semantic_key 格式: domain.type_id.xxx → 提取第二段
    sk_parts = semantic_key.split('.')
    expected_type_id = sk_parts[1] if len(sk_parts) >= 2 else ''
    if template_type_id and expected_type_id and template_type_id != expected_type_id:
        return jsonify({'success': False, 'error': f'模板类型({template_type_id})与场景对象({expected_type_id})不匹配'}), 400
    update_template_resource(template_key, {'enabled': True})
    mapping = set_semantic_default_template(semantic_key, template_key, updated_by=getattr(current_user, 'id', 'unknown'), updated_at=_x_now())
    log_action(current_user.id if current_user.is_authenticated else 'unknown', '设置语义默认模板', semantic_key, template_key)
    return jsonify({'success': True, 'semantic_key': semantic_key, 'template_key': template_key, 'mapping': mapping})




@template_mgmt_bp.route('/admin/api/template-semantic-mappings/attach', methods=['POST'])
@login_required
@require_permission('admin.templates.mapping.semantic_manage')
def admin_api_template_semantic_mapping_attach():
    data = request.get_json(silent=True) or {}
    semantic_key = str(data.get('semantic_key', '')).strip()
    template_key = str(data.get('template_key', '')).strip()
    if not semantic_key or not template_key:
        return jsonify({'success': False, 'error': 'semantic_key / template_key 不能为空'}), 400
    from template_resources import list_registered_template_resources, attach_template_key_to_semantic
    overlay = list_registered_template_resources()
    current = overlay.get(template_key)
    if not isinstance(current, dict):
        return jsonify({'success': False, 'error': '该模板尚未注册'}), 400
    # 校验 type_id 匹配
    template_type_id = current.get('type_id', '')
    sk_parts = semantic_key.split('.')
    expected_type_id = sk_parts[1] if len(sk_parts) >= 2 else ''
    if template_type_id and expected_type_id and template_type_id != expected_type_id:
        return jsonify({'success': False, 'error': f'模板类型({template_type_id})与场景对象({expected_type_id})不匹配'}), 400
    mapping = attach_template_key_to_semantic(semantic_key, template_key, updated_by=getattr(current_user, 'id', 'unknown'))
    log_action(current_user.id if current_user.is_authenticated else 'unknown', '加入语义模板候选', semantic_key, template_key)
    return jsonify({'success': True, 'semantic_key': semantic_key, 'template_key': template_key, 'mapping': mapping})




@template_mgmt_bp.route('/admin/api/templates/<template_id>')
@login_required
@require_permission('admin.templates.view')
def admin_api_template_detail(template_id):
    """模板详情 API"""
    from template_resources import TEMPLATE_MAP, list_registered_template_resources, get_type_template_mapping
    template_map = TEMPLATE_MAP
    if template_id not in template_map:
        return jsonify({'error': '模板不存在'}), 404

    info = template_map[template_id]
    dir_path = TEMPLATE_BASE / info['path']
    overlay = list_registered_template_resources()
    mapping = get_type_template_mapping(template_id)
    default_template_key = mapping.get('default_template_key', '')
    allowed_template_keys = set(mapping.get('allowed_template_keys', []) or [])
    files = []
    builtin_names = set(info.get('files', []) or [])

    for fname in info.get('files', []) or []:
        fpath = dir_path / fname
        matched_items = [
            (k, v) for k, v in overlay.items()
            if isinstance(v, dict) and v.get('type_id') == template_id and v.get('template_name') == fname
        ]
        if matched_items:
            for matched_key, matched_item in matched_items:
                files.append({
                    'name': fname,
                    'path': str(fpath),
                    'exists': fpath.exists(),
                    'size': fpath.stat().st_size if fpath.exists() else 0,
                    'template_key': matched_key,
                    'source': '注册配置',
                    'enabled': (matched_item or {}).get('enabled', True),
                    'version': (matched_item or {}).get('version', '—') if matched_item else '—',
                    'last_verified_at': (matched_item or {}).get('last_verified_at', ''),
                    'last_verify_result': _human_verify_result((matched_item or {}).get('last_verify_result', 'unverified')) if matched_item else '未做验证',
                    'is_default': matched_key == default_template_key,
                    'is_allowed': matched_key in allowed_template_keys,
                })
        elif fpath.exists():
            files.append({
                'name': fname,
                'path': str(fpath),
                'exists': True,
                'size': fpath.stat().st_size if fpath.exists() else 0,
                'template_key': '',
                'source': '内置模板',
                'enabled': True,
                'version': '—',
                'last_verified_at': '',
                'last_verify_result': '未做验证',
                'is_default': False,
                'is_allowed': False,
            })

    for matched_key, matched_item in overlay.items():
        if not isinstance(matched_item, dict) or matched_item.get('type_id') != template_id:
            continue
        template_name = str(matched_item.get('template_name', '')).strip()
        if template_name in builtin_names:
            continue
        template_path = str(matched_item.get('template_path', '')).strip()
        fpath = Path(template_path) if template_path else (dir_path / template_name if template_name else dir_path)
        files.append({
            'name': template_name or matched_key,
            'path': str(fpath),
            'exists': fpath.exists(),
            'size': fpath.stat().st_size if fpath.exists() else 0,
            'template_key': matched_key,
            'source': '注册配置',
            'enabled': matched_item.get('enabled', True),
            'version': matched_item.get('version', '—') or '—',
            'last_verified_at': matched_item.get('last_verified_at', ''),
            'last_verify_result': _human_verify_result(matched_item.get('last_verify_result', 'unverified')) or '未做验证',
            'is_default': matched_key == default_template_key,
            'is_allowed': matched_key in allowed_template_keys,
        })

    files = sorted(
        files,
        key=lambda x: (
            0 if x.get('template_key') else 1,
            str(x.get('name', '')),
            str(x.get('template_key', '')),
        )
    )

    detail_stats = {
        'all': _summarize_template_detail_rows(files),
        'main': _summarize_template_detail_rows(files, 'main') if template_id == 'operating_room' else None,
        'eye': _summarize_template_detail_rows(files, 'eye') if template_id == 'operating_room' else None,
        'aux': _summarize_template_detail_rows(files, 'aux') if template_id == 'operating_room' else None,
    }

    log_action(current_user.id if current_user.is_authenticated else 'unknown', '查看模板', template_id, info['name'])
    return jsonify({
        'id': template_id,
        'name': info.get('name', template_id),
        'domain': info.get('domain', ''),
        'path': str(dir_path),
        'files': files,
        'stats': detail_stats,
        'template_base': str(TEMPLATE_BASE),
        'mapping': mapping,
        'default_warning': (
            'missing' if default_template_key and not any((f.get('template_key') == default_template_key and f.get('exists')) for f in files)
            else 'disabled' if default_template_key and any((f.get('template_key') == default_template_key and f.get('enabled') is False) for f in files)
            else 'unset' if not default_template_key
            else ''
        ),
    })




@template_mgmt_bp.route('/api/x/template_probe', methods=['POST'])
@login_required
@require_permission('record.export')
def api_x_template_probe():
    data = request.get_json(silent=True) or {}
    project = normalize_project_payload(data.get('project', data) or {}, source='template_probe')
    validation_error = validate_normalized_project(project)
    if validation_error:
        return jsonify({'success': False, 'error': validation_error}), 400
    payload = _build_export_payload(project)
    room = payload.get('room', {}) or {}
    template_path = _x_select_template(
        payload.get('project', {}).get('domain', ''),
        room.get('level_name', '') or room.get('clean_class', ''),
        room.get('type_id', ''),
    )
    return jsonify({
        'success': True,
        'template_found': bool(template_path),
        'template_path': str(template_path) if template_path else '',
        'type_id': room.get('type_id', ''),
        'level_name': room.get('level_name', '') or room.get('clean_class', ''),
        'domain': payload.get('project', {}).get('domain', ''),
    })




@template_mgmt_bp.route('/admin/api/templates/<template_id>/upload', methods=['POST'])
@login_required
@require_permission('admin.templates.registry.manage')
def admin_api_template_upload(template_id):
    """上传/替换模板文件"""
    if not _setting_enabled('template.allow_upload', True):
        return jsonify({'success': False, 'error': '系统设置已禁止模板上传'}), 403
    from template_resources import TEMPLATE_MAP
    if template_id not in TEMPLATE_MAP:
        return jsonify({'success': False, 'error': '模板不存在'}), 404
    
    if 'file' not in request.files:
        return jsonify({'success': False, 'error': '未选择文件'}), 400
    
    file = request.files['file']
    if not file.filename:
        return jsonify({'success': False, 'error': '文件名为空'}), 400
    
    if not file.filename.endswith('.docx'):
        return jsonify({'success': False, 'error': '仅支持 .docx 格式'}), 400
    
    info = TEMPLATE_MAP[template_id]
    dir_path = TEMPLATE_BASE / info['path']
    dir_path.mkdir(parents=True, exist_ok=True)
    
    # 备份旧文件（版本历史）
    backup_dir = dir_path / '.backup'
    backup_dir.mkdir(exist_ok=True)
    
    safe_name = Path(file.filename).name
    target_path = dir_path / safe_name
    if target_path.exists():
        from datetime import datetime
        backup_name = f"{target_path.stem}_{datetime.now().strftime('%Y%m%d_%H%M%S')}{target_path.suffix}"
        import shutil
        shutil.copy2(str(target_path), str(backup_dir / backup_name))
    
    # 保存新文件
    file.save(str(target_path))
    
    # 校验文件完整性
    try:
        from docx import Document
        Document(str(target_path))
        valid = True
    except:
        valid = False
    
    log_action(current_user.id if current_user.is_authenticated else 'unknown', 
               '上传模板', template_id, f'{info["name"]} - {file.filename}')
    
    return jsonify({
        'success': True,
        'message': f'模板文件 {file.filename} 上传成功',
        'valid': valid,
        'path': str(target_path)
    })




@template_mgmt_bp.route('/admin/api/templates/<template_id>/versions')
@login_required
@require_permission('admin.templates.view')
def admin_api_template_versions(template_id):
    """模板版本历史"""
    from template_resources import TEMPLATE_MAP
    if template_id not in TEMPLATE_MAP:
        return jsonify({'success': False, 'error': '模板不存在'}), 404
    
    info = TEMPLATE_MAP[template_id]
    dir_path = TEMPLATE_BASE / info['path']
    backup_dir = dir_path / '.backup'
    
    versions = []
    if backup_dir.exists():
        for f in sorted(backup_dir.glob('*.docx'), key=lambda x: x.stat().st_mtime, reverse=True):
            versions.append({
                'name': f.name,
                'size': f.stat().st_size,
                'mtime': f.stat().st_mtime,
                'mtime_str': datetime.fromtimestamp(f.stat().st_mtime).strftime('%Y-%m-%d %H:%M:%S')
            })
    
    return jsonify({
        'template_id': template_id,
        'template_name': info['name'],
        'versions': versions,
        'total': len(versions)
    })




@template_mgmt_bp.route('/admin/api/templates/<template_id>/preview')
@login_required
@require_permission('admin.templates.preview')
def admin_api_template_preview(template_id):
    """模板预览（返回完整 HTML）"""
    from template_resources import TEMPLATE_MAP
    if template_id not in TEMPLATE_MAP:
        return jsonify({'success': False, 'error': '模板不存在'}), 404
    
    filename = request.args.get('file')
    if not filename or '..' in filename or '/' in filename or chr(92) in filename:
        return jsonify({'success': False, 'error': '非法文件名'}), 400
    
    info = TEMPLATE_MAP[template_id]
    file_path = TEMPLATE_BASE / info['path'] / filename
    
    if not file_path.exists():
        return jsonify({'success': False, 'error': '文件不存在'}), 404
    
    try:
        import mammoth
        with open(str(file_path), 'rb') as f:
            result = mammoth.convert_to_html(f)
            html_content = result.value
        
        return jsonify({
            'success': True,
            'html': html_content,
            'filename': filename,
            'file_size': file_path.stat().st_size
        })
    except Exception as e:
        return jsonify({'success': False, 'error': f'预览失败: {str(e)}'}), 500


@template_mgmt_bp.route('/admin/api/templates/<template_id>/variables')
@login_required
@require_permission('admin.templates.variables')
def admin_api_template_variables(template_id):
    """提取模板填写项（表格标签 + 关键段落关键词）"""
    from template_resources import TEMPLATE_MAP
    import re
    if template_id not in TEMPLATE_MAP:
        return jsonify({'success': False, 'error': '模板不存在'}), 404

    filename = request.args.get('file')
    if not filename or '..' in filename or '/' in filename or chr(92) in filename:
        return jsonify({'success': False, 'error': '非法文件名'}), 400

    info = TEMPLATE_MAP[template_id]
    file_path = TEMPLATE_BASE / info['path'] / filename
    if not file_path.exists():
        return jsonify({'success': False, 'error': '文件不存在'}), 404

    try:
        from docx import Document
        doc = Document(str(file_path))
        variables = []
        seen = set()

        # 1. 扫描段落：提取 {{...}} 或 【...】 占位符
        for i, para in enumerate(doc.paragraphs):
            for m in re.findall(r'\{\{([^}]+)\}\}|【([^】]+)】', para.text):
                name = m[0] or m[1]
                if name and name not in seen:
                    seen.add(name)
                    variables.append({'name': name, 'type': 'placeholder', 'location': f'段落 {i+1}'})

        # 2. 扫描表格：左列标签 → 右列为空或待填写
        for ti, table in enumerate(doc.tables):
            for ri, row in enumerate(table.rows):
                cells = [c.text.strip() for c in row.cells]
                if len(cells) < 2:
                    continue
                # 左列有内容（标签），且列数为偶数（label-value 对）
                for ci in range(0, len(cells)-1, 2):
                    label = cells[ci].strip('：: ').strip()
                    value = cells[ci+1] if ci+1 < len(cells) else ''
                    # 标签有意义（非纯数字/符号）且值为空或为占位符
                    if label and len(label) > 1 and not label.isdigit():
                        if label not in seen:
                            seen.add(label)
                            variables.append({
                                'name': label,
                                'type': 'table_field',
                                'location': f'表{ti+1} 第{ri+1}行'
                            })

        # 按类型分组输出
        return jsonify({
            'success': True,
            'variables': variables,
            'total': len(variables),
            'summary': {
                'placeholder': sum(1 for v in variables if v['type']=='placeholder'),
                'table_field': sum(1 for v in variables if v['type']=='table_field'),
            }
        })
    except Exception as e:
        return jsonify({'success': False, 'error': f'解析失败: {str(e)}'}), 500

if __name__ == '__main__':
    print(f"🚀 X1 skeleton running at http://{APP_HOST}:{APP_PORT}")
    app.run(host=APP_HOST, port=APP_PORT, debug=False)



