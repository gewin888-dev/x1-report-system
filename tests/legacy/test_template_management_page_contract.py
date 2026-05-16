#!/usr/bin/env python3
"""模板管理页面联调契约测试：详情接口与页面所需字段保持一致"""
import importlib
import json
import sys
from pathlib import Path
from tempfile import TemporaryDirectory

sys.path.insert(0, str(Path(__file__).parent))


def _login(client):
    client.get('/login')
    resp = client.post('/login', data={'username': 'admin', 'password': 'pudi2026'}, follow_redirects=True)
    return resp.status_code == 200


def main():
    print('=' * 76)
    print('模板管理页面联调契约测试')
    print('=' * 76)

    app_x1 = importlib.import_module('app_x1')
    tr = importlib.import_module('template_resources')

    with TemporaryDirectory() as td:
        temp_dir = Path(td)
        registry_path = temp_dir / 'template_registry.json'
        mappings_path = temp_dir / 'template_type_mappings.json'
        file_a = temp_dir / '制药工业兽药车间A级检测报告模板.docx'
        file_b = temp_dir / '制药工业兽药车间A级检测报告模板-候选.docx'
        from docx import Document
        d1 = Document(); d1.add_paragraph('A'); d1.save(str(file_a))
        d2 = Document(); d2.add_paragraph('B'); d2.save(str(file_b))

        registry_path.write_text(json.dumps({
            'pharma/veterinary_gmp_workshop/grade/a': {
                'type_id': 'veterinary_gmp_workshop',
                'template_name': file_a.name,
                'template_path': str(file_a),
                'resource_status': 'confirmed',
                'enabled': True,
                'version': 'v1',
                'last_verified_at': '2026-05-06 20:40:00',
                'last_verify_result': 'success'
            },
            'pharma/veterinary_gmp_workshop/grade/a-v2': {
                'type_id': 'veterinary_gmp_workshop',
                'template_name': file_b.name,
                'template_path': str(file_b),
                'resource_status': 'confirmed',
                'enabled': False,
                'version': 'v2',
                'last_verified_at': '2026-05-06 20:41:00',
                'last_verify_result': 'success'
            }
        }, ensure_ascii=False, indent=2), encoding='utf-8')
        mappings_path.write_text(json.dumps({
            'veterinary_gmp_workshop': {
                'allowed_template_keys': [
                    'pharma/veterinary_gmp_workshop/grade/a',
                    'pharma/veterinary_gmp_workshop/grade/a-v2'
                ],
                'default_template_key': 'pharma/veterinary_gmp_workshop/grade/a'
            }
        }, ensure_ascii=False, indent=2), encoding='utf-8')

        original_registry = tr.REGISTRY_FILE
        original_mappings = tr.TYPE_MAPPINGS_FILE
        tr.REGISTRY_FILE = registry_path
        tr.TYPE_MAPPINGS_FILE = mappings_path
        try:
            with app_x1.app.test_client() as client:
                assert _login(client), '登录失败'

                print('\n[CASE 1] 一级页总览接口应返回候选/默认/启用/异常汇总字段')
                resp1 = client.get('/admin/api/templates')
                assert resp1.status_code == 200
                data1 = resp1.get_json()
                item = next(t for t in data1['templates'] if t['id'] == 'veterinary_gmp_workshop')
                for field in ['default_template_key', 'candidate_count', 'enabled_count', 'missing_count', 'registered_keys']:
                    assert field in item, field
                assert item['default_template_key'] == 'pharma/veterinary_gmp_workshop/grade/a', item
                assert item['candidate_count'] == 2, item
                assert item['enabled_count'] == 1, item
                print('✅ 一级页治理汇总字段完整')

                print('\n[CASE 2] 二级页详情接口应返回 is_default / is_allowed 标志')
                resp2 = client.get('/admin/api/templates/veterinary_gmp_workshop')
                assert resp2.status_code == 200
                data2 = resp2.get_json()
                rows = {f['template_key']: f for f in data2['files'] if f.get('template_key')}
                assert rows['pharma/veterinary_gmp_workshop/grade/a']['is_default'] is True
                assert rows['pharma/veterinary_gmp_workshop/grade/a']['is_allowed'] is True
                assert rows['pharma/veterinary_gmp_workshop/grade/a-v2']['is_default'] is False
                assert rows['pharma/veterinary_gmp_workshop/grade/a-v2']['is_allowed'] is True
                print('✅ 二级页模板状态标志完整')

                print('\n[CASE 3] 前端页面脚本应包含默认模板/候选模板筛选与注册同步文案')
                html_text = Path('templates/admin.html').read_text(encoding='utf-8')
                for keyword in ['仅默认模板', '仅候选模板', '自动加入该检测类型候选模板', '注册后直接设为默认模板', '默认模板生效中', '模板治理总览', '当前风险等级']:
                    assert keyword in html_text, keyword
                print('✅ 页面脚本包含治理交互关键文案')

        finally:
            tr.REGISTRY_FILE = original_registry
            tr.TYPE_MAPPINGS_FILE = original_mappings

    print('-' * 76)
    print('✅ ALL PASS：页面联调契约成立')


if __name__ == '__main__':
    main()
