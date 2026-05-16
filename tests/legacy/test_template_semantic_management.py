#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import json
import os
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent
sys.path.insert(0, str(ROOT))

from app_x1 import app, _build_export_payload

ADMIN_HTML = (ROOT / 'templates' / 'admin.html').read_text(encoding='utf-8')
from template_resources import (
    set_type_default_template,
    set_semantic_default_template,
    get_semantic_template_mapping,
)


def assert_true(cond, msg):
    if not cond:
        raise AssertionError(msg)


def seed_overlay_template(key, type_id, name='测试模板.docx'):
    path = ROOT / 'uploads' / 'templates' / f"{key.replace('/', '_')}.docx"
    path.parent.mkdir(parents=True, exist_ok=True)
    if not path.exists() or path.stat().st_size < 1024:
        from docx import Document
        doc = Document()
        doc.add_paragraph(f'test template for {key}')
        doc.save(str(path))
    registry_path = ROOT / 'template_registry.json'
    data = {}
    if registry_path.exists():
        data = json.loads(registry_path.read_text(encoding='utf-8') or '{}')
    data[key] = {
        'template_name': name,
        'template_path': str(path),
        'type_id': type_id,
        'enabled': True,
        'version': 'v-test',
        'last_verified_at': '',
        'last_verify_result': 'smoke_success',
        'last_verify_error': '',
    }
    registry_path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding='utf-8')


def export_debug(project):
    return _build_export_payload(project)


def main():
    semantic_key = 'pharma.gmp_workshop.grade.a'
    type_id = 'gmp_workshop'
    template_a = 'pharma/gmp_workshop/grade/a'
    template_b = 'pharma/gmp_workshop/grade/b'

    seed_overlay_template(template_a, type_id, 'GMP-A模板.docx')
    seed_overlay_template(template_b, type_id, 'GMP-B模板.docx')

    set_type_default_template(type_id, template_b, updated_by='test', updated_at='2026-05-06 21:40:00')
    set_semantic_default_template(semantic_key, template_a, updated_by='test', updated_at='2026-05-06 21:40:00')

    project = {
        'project_name': 'semantic-priority-test',
        'domain': 'pharma',
        'rooms': [{
            'type_id': type_id,
            'level_name': 'A级',
            'context': {'gmp_grade': 'A级'},
            'samples': []
        }]
    }

    with app.test_client() as client:
        debug = export_debug(project)
        assert_true(debug.get('template_rule', {}).get('template_key') == template_a, f"CASE1 semantic 未优先命中: {debug}")
        print('✅ CASE1 semantic 优先于 type fallback')

        mapping = get_semantic_template_mapping(semantic_key)
        assert_true(mapping.get('default_template_key') == template_a, 'CASE2 semantic default 未写入')
        print('✅ CASE2 semantic default 可读取')

        resp = client.post('/login', data={'username': 'admin', 'password': 'pudi2026'}, follow_redirects=True)
        assert_true(resp.status_code == 200, '登录失败')

        r0 = client.get('/admin/api/template-semantic-mappings/options')
        d0 = r0.get_json() or {}
        assert_true(r0.status_code == 200 and any(x.get('semantic_key') == semantic_key for x in (d0.get('options') or [])), f'CASE0 semantic options 接口失败: {d0}')
        print('✅ CASE0 semantic options 接口可用')

        r1 = client.get('/admin/api/template-semantic-mappings/pharma.gmp_workshop.grade.a?type_id=gmp_workshop')
        d1 = r1.get_json() or {}
        item_keys = [x.get('template_key') for x in (d1.get('items') or [])]
        assert_true(r1.status_code == 200 and d1.get('mapping', {}).get('default_template_key') == template_a and template_b in item_keys, f'CASE3 detail 接口失败: {d1}')
        print('✅ CASE3 semantic detail 接口可用，且可看到同 type 注册模板')

        r2 = client.post('/admin/api/template-semantic-mappings/set-default', json={'semantic_key': semantic_key, 'template_key': template_b})
        d2 = r2.get_json() or {}
        assert_true(r2.status_code == 200 and d2.get('success') is True, f'CASE4 set-default 接口失败: {d2}')
        debug2 = export_debug(project)
        assert_true(debug2.get('template_rule', {}).get('template_key') == template_b, f"CASE4 semantic default 切换未生效: {debug2}")
        print('✅ CASE4 semantic default 切换真实影响导出链')

        import time
        template_c = f'pharma/gmp_workshop/grade/c-test-{int(time.time())}'
        # 生成有效 docx 文件，放在 TEMPLATE_BASE 内以通过路径穿越防护
        from pathlib import Path as _Path
        import os as _os
        _tb = _Path(_os.path.expanduser('~/公司资料/检测部/检测报告模板'))
        reg_file_path = _tb / 'test_tmp' / f"{template_c.replace('/', '_')}.docx"
        reg_file_path.parent.mkdir(parents=True, exist_ok=True)
        if not reg_file_path.exists() or reg_file_path.stat().st_size < 1024:
            from docx import Document
            doc = Document()
            doc.add_paragraph(f'test template for {template_c}')
            doc.save(str(reg_file_path))
        r3 = client.post('/admin/api/template-registry/register', json={
            'type_id': type_id,
            'template_key': template_c,
            'template_name': 'GMP-C模板.docx',
            'path_mode': 'absolute',
            'absolute_path': str(reg_file_path),
            'attach_to_type': False,
            'set_as_default': False,
            'semantic_key': semantic_key,
            'attach_to_semantic': True,
            'set_as_semantic_default': True,
        })
        d3 = r3.get_json() or {}
        assert_true(r3.status_code == 200 and d3.get('success') is True, f'CASE5 register 接口失败: {d3}')
        debug3 = export_debug(project)
        assert_true(debug3.get('template_rule', {}).get('template_key') == template_c, f"CASE5 注册同步到 semantic 未生效: {debug3}")
        print('✅ CASE5 注册接口可直接同步到 semantic 默认模板')

        assert_true('template-semantic-block' in ADMIN_HTML and 'renderSemanticGovernanceBlock' in ADMIN_HTML, 'CASE6 页面未接入 semantic 治理区块')
        assert_true('setSemanticDefault(' in ADMIN_HTML and 'attachSemanticCandidate(' in ADMIN_HTML and '请选择模板' in ADMIN_HTML, 'CASE7 页面未接入 semantic 下拉操作')
        assert_true('filterTemplateDetailBySemantic' in ADMIN_HTML and 'data-template-key' in ADMIN_HTML and '默认模板已停用' in ADMIN_HTML, 'CASE8 页面未接入 semantic 异常态/列表联动')
        assert_true('refreshOnly' in ADMIN_HTML and '可选模板：' in ADMIN_HTML and 'data-template-id' in ADMIN_HTML, 'CASE9 页面未接入局部刷新/候选摘要')
        assert_true('typeHasSemanticGovernance' in ADMIN_HTML and '配置方式：按业务场景分别配置' in ADMIN_HTML, 'CASE10 页面未去除语义对象的重复 type 级按钮/提示')
        print('✅ CASE6 模板详情页已接入 semantic 治理区块')
        print('✅ CASE7 模板详情页已接入 semantic 下拉操作')
        print('✅ CASE8 模板详情页已接入 semantic 异常态与列表联动')
        print('✅ CASE9 模板详情页已接入局部刷新与候选摘要')
        print('✅ CASE10 语义对象已弱化重复 type 级按钮/提示')

    # cleanup: remove test-injected keys from registry and semantic_mappings
    try:
        _reg_path = ROOT / 'template_registry.json'
        _reg = json.loads(_reg_path.read_text(encoding='utf-8') or '{}')
        _rm = [k for k in _reg if k.startswith('pharma/gmp_workshop/grade/c-test-')]
        for k in _rm:
            del _reg[k]
        if _rm:
            _reg_path.write_text(json.dumps(_reg, ensure_ascii=False, indent=2), encoding='utf-8')
            print(f'cleanup: removed {len(_rm)} test keys from registry')
    except Exception:
        pass
    try:
        _sm_path = ROOT / 'template_semantic_mappings.json'
        _sm = json.loads(_sm_path.read_text(encoding='utf-8') or '{}')
        _sm_cleaned = 0
        for _sk, _cfg in _sm.items():
            _allowed = _cfg.get('allowed_template_keys', [])
            _before = len(_allowed)
            _allowed = [k for k in _allowed if 'c-test-' not in k]
            if len(_allowed) < _before:
                _cfg['allowed_template_keys'] = _allowed
                _sm_cleaned += _before - len(_allowed)
        if _sm_cleaned:
            _sm_path.write_text(json.dumps(_sm, ensure_ascii=False, indent=2), encoding='utf-8')
            print(f'cleanup: removed {_sm_cleaned} c-test keys from semantic_mappings')
    except Exception:
        pass

    print('\n' + chr(9989) + ' ALL PASS')


if __name__ == '__main__':
    main()
