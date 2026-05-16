#!/usr/bin/env python3
"""多模板并存不串扰闭环测试：同一检测类型下多个模板项共存，停用 A 不串扰 B"""
import importlib
import json
import sys
from pathlib import Path
from tempfile import TemporaryDirectory

sys.path.insert(0, str(Path(__file__).parent))


def _login(client):
    client.get('/login')
    resp = client.post('/login', data={'username': 'admin', 'password': 'pudi2026'}, follow_redirects=True)
    return resp.status_code == 200


def _build_vet_gmp_project():
    return {
        'project_name': '多模板并存不串扰测试-兽药GMP A级',
        'report_number': 'TMPL-ISOLATION-001',
        'client_name': '测试客户',
        'contact_info': '13800000000',
        'project_address': '上海市测试区',
        'inspection_area': 'A级洁净区',
        'detection_date': '2026-05-06',
        'domain': 'pharma',
        'domain_name': '制药工业',
        'rooms': [{
            'room_id': 'r1',
            'type_id': 'veterinary_gmp_workshop',
            'room_name': 'A级车间',
            'type_name': '兽药GMP车间',
            'level_name': 'A级',
            'clean_class': 'A级',
            'basis': ['GB 50457-2019'],
            'judgement': ['GB 50457-2019'],
            'params': [],
            'summary': {
                'result_state': '合格',
                'judgement_active': ['GB 50457-2019'],
                'basis_primary': 'GB 50457-2019',
                'judgement_primary': 'GB 50457-2019'
            },
            'context': {
                'gmp_grade': 'A级'
            }
        }]
    }


def _submit_export(client, project):
    resp = client.post('/api/x/submit_export', data=json.dumps({'project': project}), content_type='application/json')
    assert resp.status_code == 200, resp.status_code
    data = resp.get_json()
    assert data.get('success') is True, data
    return data


def main():
    print('=' * 76)
    print('多模板并存不串扰闭环测试')
    print('=' * 76)

    app_x1 = importlib.import_module('app_x1')
    tr = importlib.import_module('template_resources')

    with TemporaryDirectory() as td:
        temp_dir = Path(td)
        registry_path = temp_dir / 'template_registry.json'
        from docx import Document

        file_a = temp_dir / '制药工业兽药车间A级检测报告模板.docx'
        doc_a = Document(); doc_a.add_paragraph('A 模板'); doc_a.save(str(file_a))
        file_b = temp_dir / '制药工业兽药车间A级检测报告模板-备选.docx'
        doc_b = Document(); doc_b.add_paragraph('B 模板'); doc_b.save(str(file_b))

        seed = {
            'pharma/veterinary_gmp_workshop/grade/a': {
                'type_id': 'veterinary_gmp_workshop',
                'template_name': file_a.name,
                'template_path': str(file_a),
                'resource_status': 'confirmed',
                'resource_note': 'A 主模板',
                'enabled': True,
                'version': 'v1',
                'last_verified_at': '2026-05-06 20:20:00',
                'last_verify_result': 'success',
                'last_verify_error': ''
            },
            'pharma/veterinary_gmp_workshop/grade/a-v2': {
                'type_id': 'veterinary_gmp_workshop',
                'template_name': file_b.name,
                'template_path': str(file_b),
                'resource_status': 'confirmed',
                'resource_note': 'A 备选模板',
                'enabled': True,
                'version': 'v2',
                'last_verified_at': '2026-05-06 20:21:00',
                'last_verify_result': 'success',
                'last_verify_error': ''
            }
        }
        registry_path.write_text(json.dumps(seed, ensure_ascii=False, indent=2), encoding='utf-8')

        original_registry = tr.REGISTRY_FILE
        tr.REGISTRY_FILE = registry_path
        try:
            project = _build_vet_gmp_project()
            with app_x1.app.test_client() as client:
                assert _login(client), '登录失败'

                print('\n[CASE 1] 详情页应同时返回两个注册模板项')
                detail_resp = client.get('/admin/api/templates/veterinary_gmp_workshop')
                assert detail_resp.status_code == 200
                detail = detail_resp.get_json()
                keys = [f.get('template_key') for f in detail.get('files', []) if f.get('template_key')]
                assert 'pharma/veterinary_gmp_workshop/grade/a' in keys
                assert 'pharma/veterinary_gmp_workshop/grade/a-v2' in keys
                print('✅ 详情页双模板共存可见')

                print('\n[CASE 2] 导出链当前只命中规则 key=A')
                data1 = _submit_export(client, project)
                tr1 = data1['export_payload'].get('template_resource', {})
                assert tr1.get('template_key') == 'pharma/veterinary_gmp_workshop/grade/a', tr1
                assert tr1.get('template_path') == str(file_a), tr1
                assert tr1.get('template_found') is True, tr1
                print('✅ 导出链只命中 A，未误命中 B')

                print('\n[CASE 3] 停用 A 后，不应串扰 B；详情页 B 仍启用')
                toggle_resp = client.post('/admin/api/template-registry/toggle', json={
                    'template_key': 'pharma/veterinary_gmp_workshop/grade/a',
                    'enabled': False
                })
                assert toggle_resp.status_code == 200
                toggle_data = toggle_resp.get_json()
                assert toggle_data.get('enabled') is False

                after = json.loads(registry_path.read_text(encoding='utf-8'))
                assert after['pharma/veterinary_gmp_workshop/grade/a']['enabled'] is False
                assert after['pharma/veterinary_gmp_workshop/grade/a-v2']['enabled'] is True

                detail_resp2 = client.get('/admin/api/templates/veterinary_gmp_workshop')
                detail2 = detail_resp2.get_json()
                rows = {f.get('template_key'): f for f in detail2.get('files', []) if f.get('template_key')}
                assert rows['pharma/veterinary_gmp_workshop/grade/a']['enabled'] is False
                assert rows['pharma/veterinary_gmp_workshop/grade/a-v2']['enabled'] is True
                print('✅ 停用 A 未串扰 B，详情页状态一致')

                print('\n[CASE 4] 停用 A 后，导出链仅对 A 感知 disabled，不会错误切到 B')
                data2 = _submit_export(client, project)
                tr2 = data2['export_payload'].get('template_resource', {})
                assert tr2.get('template_key') == 'pharma/veterinary_gmp_workshop/grade/a', tr2
                assert tr2.get('resource_status') == 'disabled', tr2
                assert tr2.get('template_found') is False, tr2
                print('✅ 导出链仅感知 A disabled，未被 B 串扰')

        finally:
            tr.REGISTRY_FILE = original_registry

    print('-' * 76)
    print('✅ ALL PASS：同类型多模板并存时，治理动作不串扰')


if __name__ == '__main__':
    main()
