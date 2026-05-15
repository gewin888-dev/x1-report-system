#!/usr/bin/env python3
"""默认模板切换闭环测试：检测类型默认模板切换后，导出链真实切换命中结果"""
import importlib
import json
import sys
from pathlib import Path
from tempfile import TemporaryDirectory

sys.path.insert(0, str(Path(__file__).parent))


def _login(client):
    client.get('/login')
    resp = client.post('/login', data={'username': 'admin', 'password': 'pudi2026'}, follow_redirects=True)
    return resp.status_code == 200


def _build_vet_gmp_project():
    return {
        'project_name': '默认模板切换测试-兽药GMP A级',
        'report_number': 'TMPL-DEFAULT-SWITCH-001',
        'client_name': '测试客户',
        'contact_info': '13800000000',
        'project_address': '上海市测试区',
        'inspection_area': 'A级洁净区',
        'detection_date': '2026-05-06',
        'domain': 'pharma',
        'domain_name': '制药工业',
        'rooms': [{
            'room_id': 'r1',
            'type_id': 'veterinary_gmp_workshop',
            'room_name': 'A级车间',
            'type_name': '兽药GMP车间',
            'level_name': 'A级',
            'clean_class': 'A级',
            'basis': ['GB 50457-2019'],
            'judgement': ['GB 50457-2019'],
            'params': [],
            'summary': {'result_state': '合格'},
            'context': {'gmp_grade': 'A级'}
        }]
    }


def _submit_export(client, project):
    resp = client.post('/api/x/submit_export', data=json.dumps({'project': project}), content_type='application/json')
    assert resp.status_code == 200, resp.status_code
    data = resp.get_json()
    assert data.get('success') is True, data
    return data


def main():
    print('=' * 76)
    print('默认模板切换闭环测试')
    print('=' * 76)

    app_x1 = importlib.import_module('app_x1')
    tr = importlib.import_module('template_resources')

    with TemporaryDirectory() as td:
        temp_dir = Path(td)
        registry_path = temp_dir / 'template_registry.json'
        mappings_path = temp_dir / 'template_type_mappings.json'
        from docx import Document

        file_a = temp_dir / '制药工业兽药车间A级检测报告模板.docx'
        doc_a = Document(); doc_a.add_paragraph('A 模板'); doc_a.save(str(file_a))
        file_b = temp_dir / '制药工业兽药车间A级检测报告模板-新版.docx'
        doc_b = Document(); doc_b.add_paragraph('B 模板'); doc_b.save(str(file_b))

        seed = {
            'pharma/veterinary_gmp_workshop/grade/a': {
                'type_id': 'veterinary_gmp_workshop',
                'template_name': file_a.name,
                'template_path': str(file_a),
                'resource_status': 'confirmed',
                'enabled': True,
                'version': 'v1',
                'last_verify_result': 'smoke_success',
                'last_verify_error': ''
            },
            'pharma/veterinary_gmp_workshop/grade/a-v2': {
                'type_id': 'veterinary_gmp_workshop',
                'template_name': file_b.name,
                'template_path': str(file_b),
                'resource_status': 'confirmed',
                'enabled': True,
                'version': 'v2',
                'last_verify_result': 'smoke_success',
                'last_verify_error': ''
            }
        }
        mappings = {
            'veterinary_gmp_workshop': {
                'allowed_template_keys': [
                    'pharma/veterinary_gmp_workshop/grade/a',
                    'pharma/veterinary_gmp_workshop/grade/a-v2'
                ],
                'default_template_key': 'pharma/veterinary_gmp_workshop/grade/a',
                'updated_at': '2026-05-06 20:31:00',
                'updated_by': 'seed'
            }
        }
        registry_path.write_text(json.dumps(seed, ensure_ascii=False, indent=2), encoding='utf-8')
        mappings_path.write_text(json.dumps(mappings, ensure_ascii=False, indent=2), encoding='utf-8')

        semantic_mappings_path = temp_dir / 'template_semantic_mappings.json'
        semantic_mappings_path.write_text(json.dumps({}, ensure_ascii=False, indent=2), encoding='utf-8')

        original_registry = tr.REGISTRY_FILE
        original_mappings = tr.TYPE_MAPPINGS_FILE
        original_semantic_mappings = tr.SEMANTIC_MAPPINGS_FILE
        tr.REGISTRY_FILE = registry_path
        tr.TYPE_MAPPINGS_FILE = mappings_path
        tr.SEMANTIC_MAPPINGS_FILE = semantic_mappings_path
        try:
            project = _build_vet_gmp_project()
            with app_x1.app.test_client() as client:
                assert _login(client), '登录失败'

                print('\n[CASE 1] 初始默认模板=A，导出链应命中 A')
                data1 = _submit_export(client, project)
                rule1 = data1['export_payload'].get('template_rule', {})
                res1 = data1['export_payload'].get('template_resource', {})
                assert rule1.get('template_key') == 'pharma/veterinary_gmp_workshop/grade/a', rule1
                assert res1.get('template_path') == str(file_a), res1
                print('✅ 初始默认模板 A 生效')

                print('\n[CASE 2] 设默认模板=B，映射文件应同步更新')
                set_resp = client.post('/admin/api/template-type-mappings/set-default', json={
                    'type_id': 'veterinary_gmp_workshop',
                    'template_key': 'pharma/veterinary_gmp_workshop/grade/a-v2'
                })
                assert set_resp.status_code == 200, set_resp.status_code
                set_data = set_resp.get_json()
                assert set_data.get('success') is True, set_data
                current_map = json.loads(mappings_path.read_text(encoding='utf-8'))
                assert current_map['veterinary_gmp_workshop']['default_template_key'] == 'pharma/veterinary_gmp_workshop/grade/a-v2'
                print('✅ 默认模板映射已切到 B')

                print('\n[CASE 3] 切换后再次导出，应真实命中 B')
                data2 = _submit_export(client, project)
                rule2 = data2['export_payload'].get('template_rule', {})
                res2 = data2['export_payload'].get('template_resource', {})
                assert rule2.get('template_key') == 'pharma/veterinary_gmp_workshop/grade/a-v2', rule2
                assert rule2.get('resolver') == 'x1-template-rule+type-default', rule2
                assert res2.get('template_path') == str(file_b), res2
                assert res2.get('template_found') is True, res2
                print('✅ 默认模板切换后导出链真实切到 B')

        finally:
            tr.REGISTRY_FILE = original_registry
            tr.TYPE_MAPPINGS_FILE = original_mappings
            tr.SEMANTIC_MAPPINGS_FILE = original_semantic_mappings

    print('-' * 76)
    print('✅ ALL PASS：默认模板切换已真实影响导出链')


if __name__ == '__main__':
    main()
